import os
import logging
from typing import Dict, Any, Optional, List, Tuple, TYPE_CHECKING
from pathlib import Path
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
import re
from dataclasses import dataclass
import json

# Type checking imports
if TYPE_CHECKING:
    from PIL import Image

# Optional imports for .doc and OCR support
try:
    from doc2txt import extract_text as doc2txt_extract
    DOC2TXT_AVAILABLE = True
except ImportError:
    DOC2TXT_AVAILABLE = False
    doc2txt_extract = None

try:
    import pytesseract
    from PIL import Image
    import cv2
    import io
    import numpy as np
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    np = None
    Image = None

@dataclass
class DocumentMetadata:
    """Metadata for extracted documents."""
    filename: str
    filepath: str
    file_size: int
    page_count: int
    character_count: int
    word_count: int
    extraction_method: str
    confidence_score: float
    processing_time: float
    checksum: str

@dataclass
class ExtractionResult:
    """Result of document extraction."""
    text: str
    metadata: DocumentMetadata
    warnings: List[str] = None

    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []

class AdvancedDocumentProcessor:
    """
    World-class document processor optimized for petrophysical reports.
    Handles PDFs, DOCX, and text files with maximum text extraction accuracy.
    """

    def __init__(self, max_workers: int = 2, enable_validation: bool = True):
        """
        Initialize the document processor.

        Args:
            max_workers: Number of parallel processing threads
            enable_validation: Whether to validate extraction quality
        """
        self.max_workers = max_workers
        self.enable_validation = enable_validation

        # Setup logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)

        # Initialize NLTK for sentence splitting
        try:
            import nltk
            # Try punkt_tab first (NLTK 3.9+)
            try:
                nltk.data.find('tokenizers/punkt_tab')
            except LookupError:
                # Fall back to punkt for older NLTK versions
                try:
                    nltk.data.find('tokenizers/punkt')
                except LookupError:
                    # Download punkt_tab (primary) and punkt (fallback)
                    nltk.download('punkt_tab', quiet=True)
                    try:
                        nltk.download('punkt', quiet=True)
                    except:
                        pass  # punkt may not be needed if punkt_tab works
        except Exception as e:
            self.logger.warning(f"NLTK initialization issue: {e}")

        # Check for doc2txt availability for .doc files (includes antiword binaries)
        self.doc2txt_available = DOC2TXT_AVAILABLE
        
        if not self.doc2txt_available:
            self.logger.warning("doc2txt not installed - .doc files will be skipped")
            self.logger.info("Install with: pip install doc2txt")
        else:
            self.logger.info("[OK] doc2txt available - .doc file support enabled (includes antiword binaries)")

        # Initialize OCR capabilities
        self.ocr_available = OCR_AVAILABLE
        self.ocr_language = 'eng'
        
        if not self.ocr_available:
            self.logger.warning("OCR libraries not available - OCR fallback disabled")
        else:
            # Verify Tesseract is installed
            try:
                pytesseract.get_tesseract_version()
                self.logger.info("[OK] OCR (Tesseract) available - will be used for scanned PDFs")
            except Exception as e:
                self.logger.warning(f"Tesseract OCR not available: {e}")
                self.ocr_available = False


    def process_documents(self, input_dir: str, output_dir: str = None) -> Dict[str, ExtractionResult]:
        """
        Process all documents in a directory.

        Args:
            input_dir: Directory containing documents
            output_dir: Directory to save extracted text (optional)

        Returns:
            Dictionary mapping filenames to extraction results
        """
        input_path = Path(input_dir)
        if not input_path.exists():
            raise FileNotFoundError(f"Input directory not found: {input_dir}")

        # Find all supported documents
        supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md', '.dat'}
        document_files = []

        for ext in supported_extensions:
            document_files.extend(input_path.glob(f'**/*{ext}'))
        
        # Filter out Readme.md files (they're usually not useful for RAG)
        document_files = [f for f in document_files if not f.name.lower() in ['readme.md', 'readme.txt']]

        self.logger.info(f"Found {len(document_files)} documents to process")

        # Process documents (parallel processing for performance)
        results = {}

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_file = {
                executor.submit(self._process_single_document, file_path): file_path
                for file_path in document_files
            }

            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    result = future.result()
                    if result:
                        # Use full path as key to handle duplicate filenames in different directories
                        results[str(file_path)] = result
                        self.logger.info(f"[OK] Processed: {file_path}")
                    else:
                        self.logger.warning(f"[FAILED] Failed: {file_path}")
                except Exception as e:
                    self.logger.error(f"[ERROR] Error processing {file_path}: {e}")

        # Save results if output directory provided
        if output_dir:
            self._save_results(results, output_dir)

        return results

    def _process_single_document(self, filepath: Path) -> Optional[ExtractionResult]:
        """Process a single document with comprehensive extraction."""
        import time
        start_time = time.time()

        try:
            # Calculate file checksum for change detection
            checksum = self._calculate_checksum(filepath)

            # Route to appropriate processor based on file type
            ext = filepath.suffix.lower()

            if ext == '.pdf':
                text, metadata = self._extract_pdf_comprehensive(filepath)
            elif ext == '.docx':
                text, metadata = self._extract_docx_comprehensive(filepath)
            elif ext == '.doc':
                text, metadata = self._extract_doc_file(filepath)
            elif ext in ['.txt', '.md']:
                text, metadata = self._extract_text_comprehensive(filepath)
            elif ext == '.dat':
                text, metadata = self._extract_dat_file(filepath)
            else:
                return None

            # Create metadata
            processing_time = time.time() - start_time
            doc_metadata = DocumentMetadata(
                filename=filepath.name,
                filepath=str(filepath),
                file_size=filepath.stat().st_size,
                page_count=metadata.get('page_count', 1),
                character_count=len(text),
                word_count=len(text.split()),
                extraction_method=metadata.get('method', 'unknown'),
                confidence_score=metadata.get('confidence', 0.0),
                processing_time=processing_time,
                checksum=checksum
            )

            # Validate extraction quality
            warnings = []
            if self.enable_validation:
                warnings = self._validate_extraction(text, metadata)

            return ExtractionResult(
                text=text,
                metadata=doc_metadata,
                warnings=warnings
            )

        except Exception as e:
            self.logger.error(f"Error processing {filepath}: {e}")
            return None

    def _extract_pdf_comprehensive(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using multiple methods for maximum accuracy.
        Optimized for petrophysical reports with tables and complex layouts.
        """
        methods_tried = []
        best_result = {'text': '', 'confidence': 0.0, 'method': 'none', 'page_count': 0}

        # Method 1: PyMuPDF (fitz) - Best for complex layouts
        try:
            text, stats = self._extract_pdf_fitz(filepath)
            confidence = self._calculate_pdf_confidence(text, stats)
            methods_tried.append('PyMuPDF')

            if confidence > best_result['confidence']:
                best_result = {
                    'text': text,
                    'confidence': confidence,
                    'method': 'PyMuPDF',
                    'page_count': stats.get('pages', 0)
                }
        except Exception as e:
            self.logger.debug(f"PyMuPDF failed: {e}")

        # Method 2: pdfplumber - Excellent for tables and structured data
        try:
            text, stats = self._extract_pdf_plumber(filepath)
            confidence = self._calculate_pdf_confidence(text, stats)
            methods_tried.append('pdfplumber')

            if confidence > best_result['confidence']:
                best_result = {
                    'text': text,
                    'confidence': confidence,
                    'method': 'pdfplumber',
                    'page_count': stats.get('pages', 0)
                }
        except Exception as e:
            self.logger.debug(f"pdfplumber failed: {e}")

        # Method 3: PyPDF2 as fallback
        if best_result['confidence'] < 0.7:
            try:
                text, stats = self._extract_pdf_pypdf2(filepath)
                confidence = self._calculate_pdf_confidence(text, stats)
                methods_tried.append('PyPDF2')

                if confidence > best_result['confidence']:
                    best_result = {
                        'text': text,
                        'confidence': confidence,
                        'method': 'PyPDF2',
                        'page_count': stats.get('pages', 0)
                    }
            except Exception as e:
                self.logger.debug(f"PyPDF2 failed: {e}")

        # Method 4: OCR fallback for scanned/image-based PDFs
        # Trigger OCR if confidence is low or character count is suspiciously low
        if (best_result['confidence'] < 0.5 or 
            (best_result.get('page_count', 0) > 0 and 
             len(best_result['text']) / max(best_result.get('page_count', 1), 1) < 100)):
            
            if self.ocr_available:
                try:
                    text, stats = self._extract_pdf_ocr(filepath)
                    confidence = self._calculate_pdf_confidence(text, stats)
                    methods_tried.append('OCR')

                    # Use OCR if it provides better results
                    if confidence > best_result['confidence'] or len(text) > len(best_result['text']) * 1.5:
                        best_result = {
                            'text': text,
                            'confidence': confidence,
                            'method': 'OCR',
                            'page_count': stats.get('pages', 0)
                        }
                        self.logger.info(f"OCR improved extraction for {filepath.name}")
                except Exception as e:
                    self.logger.debug(f"OCR failed: {e}")
            else:
                self.logger.debug("OCR not available for fallback")

        best_result['methods_tried'] = methods_tried
        return best_result['text'], best_result

    def _extract_pdf_fitz(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract using PyMuPDF (fitz) - best for complex layouts."""
        doc = fitz.open(str(filepath))
        text = ""
        stats = {'pages': len(doc), 'characters': 0, 'blocks': 0}

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # Try different extraction modes for maximum text capture
            page_text = page.get_text()

            # If basic extraction is poor, try block-level extraction
            if len(page_text.strip()) < 100:  # Arbitrary threshold
                blocks = page.get_text("dict")
                page_text = self._extract_text_from_fitz_blocks(blocks)

            text += page_text + "\n"
            stats['blocks'] += 1

        stats['characters'] = len(text)
        doc.close()

        return text.strip(), stats

    def _extract_pdf_plumber(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract using pdfplumber - excellent for tables."""
        text = ""
        stats = {'pages': 0, 'characters': 0, 'tables': 0}

        with pdfplumber.open(str(filepath)) as pdf:
            stats['pages'] = len(pdf.pages)

            for page in pdf.pages:
                # Extract main text
                page_text = page.extract_text() or ""
                text += page_text + "\n"

                # Extract tables and convert to text
                tables = page.extract_tables()
                for table in tables:
                    stats['tables'] += 1
                    # Convert table to readable text
                    table_text = self._table_to_text(table)
                    text += "\n[TABLE]\n" + table_text + "\n[/TABLE]\n"

        stats['characters'] = len(text)
        return text.strip(), stats

    def _extract_pdf_pypdf2(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract using PyPDF2 - reliable fallback."""
        import PyPDF2

        text = ""
        with open(str(filepath), 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            stats = {'pages': len(pdf_reader.pages), 'characters': 0}

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                text += page_text + "\n"

        stats['characters'] = len(text)
        return text.strip(), stats

    def _extract_text_from_fitz_blocks(self, blocks: Dict) -> str:
        """Extract text from PyMuPDF block dictionary."""
        text = ""
        for block in blocks.get('blocks', []):
            if 'lines' in block:
                for line in block['lines']:
                    for span in line.get('spans', []):
                        if 'text' in span:
                            text += span['text'] + " "
                    text += "\n"
        return text

    def _table_to_text(self, table: List[List]) -> str:
        """Convert table data to readable text format."""
        if not table:
            return ""

        text_lines = []
        for row in table:
            if row:
                # Clean and join row data
                clean_row = [str(cell).strip() if cell is not None else "" for cell in row]
                text_lines.append(" | ".join(clean_row))

        return "\n".join(text_lines)

    def _extract_pdf_ocr(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using OCR for scanned/image-based documents."""
        if not self.ocr_available:
            raise RuntimeError("OCR not available")
        
        doc = fitz.open(str(filepath))
        text = ""
        stats = {'pages': len(doc), 'characters': 0, 'ocr_pages': 0}
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # Convert page to image (300 DPI for quality)
            mat = fitz.Matrix(300/72, 300/72)  # 300 DPI
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Preprocess image for better OCR
            img_processed = self._preprocess_image_for_ocr(img)
            
            # Run OCR
            try:
                page_text = pytesseract.image_to_string(img_processed, lang=self.ocr_language)
                text += page_text + "\n"
                stats['ocr_pages'] += 1
            except Exception as e:
                self.logger.debug(f"OCR failed for page {page_num}: {e}")
        
        doc.close()
        stats['characters'] = len(text)
        
        return text.strip(), stats

    def _preprocess_image_for_ocr(self, img: 'Image.Image') -> 'Image.Image':
        """Preprocess image to improve OCR accuracy."""
        if not OCR_AVAILABLE:
            return img
        
        try:
            # Convert to numpy array for OpenCV processing
            img_array = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
            
            # Convert to grayscale
            gray = cv2.cvtColor(img_array, cv2.COLOR_BGR2GRAY)
            
            # Denoise
            denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
            
            # Enhance contrast using CLAHE (Contrast Limited Adaptive Histogram Equalization)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(denoised)
            
            # Convert back to PIL Image
            img_processed = Image.fromarray(enhanced)
            return img_processed
            
        except Exception as e:
            self.logger.debug(f"Image preprocessing failed: {e}, using original image")
            return img.convert('L')  # Fallback to grayscale

    def _detect_if_scanned(self, filepath: Path) -> bool:
        """Detect if PDF appears to be scanned/image-based."""
        try:
            doc = fitz.open(str(filepath))
            if len(doc) == 0:
                return False
            
            # Check first page
            page = doc.load_page(0)
            text = page.get_text()
            
            # If very little text, likely scanned
            if len(text.strip()) < 50:
                return True
            
            # Check if text extraction confidence is low
            # This is a heuristic - scanned PDFs typically have minimal extractable text
            doc.close()
            return False
            
        except Exception:
            return False

    def _extract_doc_file(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from legacy .doc files using doc2txt (includes antiword binaries)."""
        if not self.doc2txt_available:
            raise RuntimeError(
                "doc2txt not installed - cannot process .doc files. "
                "Install with: pip install doc2txt"
            )
        
        try:
            # Use doc2txt to extract text from .doc file
            # doc2txt includes antiword binaries, no separate installation needed
            text = doc2txt_extract(str(filepath))
            
            if not text or not text.strip():
                raise RuntimeError("doc2txt returned empty text")
            
            stats = {
                'paragraphs': len(text.split('\n\n')),
                'characters': len(text),
                'words': len(text.split())
            }
            
            return text, {
                'method': 'doc2txt',
                'confidence': 0.90,  # doc2txt with bundled antiword is reliable
                'page_count': 1,  # .doc files don't have clear page boundaries
                **stats
            }
            
        except Exception as e:
            error_msg = str(e)
            raise RuntimeError(f"Failed to extract text from .doc file using doc2txt: {error_msg}")

    def _extract_docx_comprehensive(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Word documents comprehensively."""
        doc = Document(str(filepath))
        text_parts = []
        stats = {'paragraphs': 0, 'characters': 0, 'tables': 0}

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                stats['paragraphs'] += 1

        # Extract tables
        for table in doc.tables:
            stats['tables'] += 1
            table_text = self._table_to_text([
                [cell.text for cell in row.cells] for row in table.rows
            ])
            text_parts.append(f"[TABLE]\n{table_text}\n[/TABLE]")

        # Handle headers/footers if available
        try:
            for section in doc.sections:
                header_text = self._extract_section_text(section.header)
                if header_text:
                    text_parts.insert(0, f"[HEADER]\n{header_text}\n[/HEADER]")

                footer_text = self._extract_section_text(section.footer)
                if footer_text:
                    text_parts.append(f"[FOOTER]\n{footer_text}\n[/FOOTER]")
        except:
            pass  # Headers/footers might not be accessible

        full_text = "\n".join(text_parts)
        stats['characters'] = len(full_text)

        return full_text, {
            'method': 'python-docx',
            'confidence': 0.95,
            'page_count': 1,  # DOCX doesn't have pages
            **stats
        }

    def _extract_section_text(self, section) -> str:
        """Extract text from document section."""
        if not section:
            return ""
        return "\n".join([para.text for para in section.paragraphs if para.text.strip()])

    def _extract_text_comprehensive(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract from plain text files."""
        with open(str(filepath), 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

        stats = {
            'lines': len(text.split('\n')),
            'characters': len(text),
            'words': len(text.split())
        }

        return text, {
            'method': 'direct_read',
            'confidence': 1.0,
            'page_count': 1,
            **stats
        }

    def _extract_dat_file(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from .dat files, with special handling for Well_picks_Volve_v1.dat.
        
        Args:
            filepath: Path to .dat file
            
        Returns:
            Tuple of (text, metadata)
        """
        import re
        
        # Check if this is the Well Picks file
        is_well_picks = 'Well_picks' in filepath.name or 'well_picks' in filepath.name.lower()
        
        if is_well_picks:
            return self._extract_well_picks_dat(filepath)
        else:
            # Generic .dat file - treat as text
            return self._extract_text_comprehensive(filepath)
    
    def _extract_well_picks_dat(self, filepath: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract and format formation picks data from Well_picks_Volve_v1.dat.
        
        Converts structured formation data into natural language text for RAG indexing.
        
        Args:
            filepath: Path to Well_picks_Volve_v1.dat file
            
        Returns:
            Tuple of (formatted_text, metadata)
        """
        import re
        
        try:
            with open(str(filepath), 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()
            
            # Parse the file
            well_data = {}
            current_well = None
            in_data_section = False
            header_line = None
            
            for line in lines:
                line = line.rstrip()
                
                # Skip comments and empty lines
                if not line or line.startswith('#'):
                    continue
                
                # Detect well header: "Well NO 15/9-11" or "Well NO 15/9-19 A" or "Well NO 15/9-A-15" or "Well NO 15/9-F-12 pilot"
                # Pattern must match: Well NO 15/9-11, Well NO 15/9-A-15, Well NO 15/9-F-12 pilot, etc.
                # But NOT match: "Well name" (column header)
                if line.strip().startswith('Well NO ') and 'Well name' not in line:
                    # Extract well name: everything after "Well "
                    well_name = line.strip()[5:].strip()  # Remove "Well " prefix
                    current_well = well_name
                    well_data[current_well] = []
                    in_data_section = False
                    header_line = None
                    continue
                
                # Detect column header line (contains "Well name" and "Surface name")
                if 'Well name' in line and 'Surface name' in line:
                    header_line = line
                    in_data_section = True
                    continue
                
                # Skip separator lines (dashes)
                if re.match(r'^[\s\-]+$', line):
                    continue
                
                # Parse data rows (only if we're in a data section and have a current well)
                if in_data_section and current_well and header_line:
                    # The format is fixed-width. Based on the header, approximate column positions:
                    # Well name: ~0-25, Surface name: ~25-65, Obs#: ~65-70, Qlf: ~70-73, MD: ~73-82, TVD: ~82-91, TVDSS: ~91-100
                    # But we'll use a more flexible approach: find well name pattern, then extract fields
                    
                    # Check if line starts with well name pattern (NO 15/9-XX or NO 15/9-XX A)
                    # Pattern: NO 15/9-11 or NO 15/9-19 A or NO 15/5-7 A
                    well_name_match = re.match(r'^\s*(NO\s+\d+[/\-]\d+[/\-][A-Z]?\d*[A-Z]*)\s+', line)
                    if well_name_match:
                        # Extract well name
                        matched_well = well_name_match.group(1).strip()
                        
                        # Check if there's a suffix like "A", "B", "SR", "BT2" right after
                        remaining_after_match = line[well_name_match.end():]
                        suffix_match = re.match(r'^([A-Z0-9]+)\s+', remaining_after_match)
                        if suffix_match and suffix_match.group(1) in ['A', 'B', 'C', 'SR', 'BT2', 'BT']:
                            # This is part of the well name, not the formation name
                            well_name_end_pos = well_name_match.end() + len(suffix_match.group(0))
                        else:
                            well_name_end_pos = well_name_match.end()
                        
                        # Extract surface name - it's between well name and Obs#
                        # Obs# is typically a single digit number, find it
                        remaining_line = line[well_name_end_pos:].strip()
                        
                        # Try to find Obs# (single digit) and Qlf (2-letter code or empty)
                        # Pattern: surface name, then Obs# (digit), then Qlf (2-3 letters or empty), then MD (number)
                        obs_match = re.search(r'(\d+)\s+([A-Z]{2,3}|\s{2,3})\s+([\d\.]+)', remaining_line)
                        if obs_match:
                            obs_num = obs_match.group(1)
                            qlf = obs_match.group(2).strip() if obs_match.group(2).strip() else None
                            md_str = obs_match.group(3)
                            
                            # Surface name is everything before Obs#
                            surface_name = remaining_line[:obs_match.start()].strip()
                            
                            # Extract MD, TVD, TVDSS from the rest of the line
                            # They're typically the next numeric values
                            after_md = remaining_line[obs_match.end():].strip()
                            
                            # Extract all numeric values (MD, TVD, TVDSS, TWT, etc.)
                            numbers = re.findall(r'[\d\.\-]+', after_md)
                            
                            md = None
                            tvd = None
                            tvds = None
                            
                            if md_str:
                                try:
                                    md = float(md_str)
                                except ValueError:
                                    pass
                            
                            # TVD and TVDSS are typically the next two numbers after MD
                            num_index = 0
                            for num_str in numbers:
                                try:
                                    val = float(num_str)
                                    if num_index == 0 and md is None:
                                        md = val
                                    elif num_index == 0 and md is not None:
                                        tvd = val
                                    elif num_index == 1 and tvd is not None:
                                        tvds = val
                                        break
                                    num_index += 1
                                except ValueError:
                                    continue
                            
                            # Only add if we have at least a formation name
                            if surface_name:
                                well_data[current_well].append({
                                    'formation': surface_name,
                                    'md': md,
                                    'tvd': tvd,
                                    'tvds': tvds,
                                    'qlf': qlf
                                })
                        else:
                            # Fallback: try simpler parsing - split by whitespace
                            parts = remaining_line.split()
                            if parts:
                                # First part(s) should be surface name
                                # Find where numbers start
                                surface_parts = []
                                num_start_idx = None
                                for i, part in enumerate(parts):
                                    try:
                                        float(part)
                                        num_start_idx = i
                                        break
                                    except ValueError:
                                        surface_parts.append(part)
                                
                                if surface_parts:
                                    surface_name = ' '.join(surface_parts).strip()
                                    
                                    # Extract numeric values
                                    if num_start_idx is not None:
                                        numbers = []
                                        for i in range(num_start_idx, min(num_start_idx + 5, len(parts))):
                                            try:
                                                numbers.append(float(parts[i]))
                                            except (ValueError, IndexError):
                                                break
                                        
                                        if len(numbers) >= 1:
                                            md = numbers[0]
                                        if len(numbers) >= 2:
                                            tvd = numbers[1]
                                        if len(numbers) >= 3:
                                            tvds = numbers[2]
                                    
                                    well_data[current_well].append({
                                        'formation': surface_name,
                                        'md': md,
                                        'tvd': tvd,
                                        'tvds': tvds,
                                        'qlf': None
                                    })
            
            # Format the data into markdown tables for better LLM parsing
            formatted_lines = []
            formatted_lines.append("Formation Picks Data from Well_picks_Volve_v1.dat")
            formatted_lines.append("=" * 60)
            formatted_lines.append("")
            
            for well_name, formations in sorted(well_data.items()):
                if not formations:
                    continue
                
                formatted_lines.append(f"Well {well_name} Formation Intervals:")
                formatted_lines.append("")
                
                # Create markdown table
                # Extract formation type (Top/Base) from formation name
                table_rows = []
                table_rows.append("| Formation Name | Type | MD (m) | TVD (m) | TVDSS (m) | Quality |")
                table_rows.append("|----------------|------|--------|---------|-----------|---------|")
                
                for form in formations:
                    formation_name = form['formation']
                    
                    # Extract type (Top/Base) from formation name
                    formation_type = "-"
                    if " Top" in formation_name:
                        formation_type = "Top"
                        formation_name = formation_name.replace(" Top", "").strip()
                    elif " Base" in formation_name:
                        formation_type = "Base"
                        formation_name = formation_name.replace(" Base", "").strip()
                    
                    # Format depth values
                    md_str = f"{form['md']:.2f}" if form['md'] is not None else "-"
                    tvd_str = f"{form['tvd']:.2f}" if form['tvd'] is not None else "-"
                    tvds_str = f"{form['tvds']:.2f}" if form['tvds'] is not None else "-"
                    
                    # Quality flag
                    if form['qlf']:
                        qlf_meanings = {
                            'ER': 'Eroded',
                            'FP': 'Faulted pick',
                            'FO': 'Faulted out',
                            'NL': 'Not logged',
                            'NR': 'Not reached'
                        }
                        quality = qlf_meanings.get(form['qlf'], form['qlf'])
                    else:
                        quality = "Not logged"
                    
                    # Create table row
                    table_rows.append(f"| {formation_name} | {formation_type} | {md_str} | {tvd_str} | {tvds_str} | {quality} |")
                
                formatted_lines.extend(table_rows)
                formatted_lines.append("")  # Empty line between wells
            
            formatted_text = "\n".join(formatted_lines)
            
            # Calculate statistics
            total_wells = len(well_data)
            total_formations = sum(len(forms) for forms in well_data.values())
            
            return formatted_text, {
                'method': 'well_picks_parser',
                'confidence': 1.0,
                'page_count': 1,
                'wells_count': total_wells,
                'formations_count': total_formations,
                'lines': len(formatted_lines),
                'characters': len(formatted_text),
                'words': len(formatted_text.split())
            }
            
        except Exception as e:
            self.logger.error(f"Error parsing Well Picks file {filepath}: {e}")
            # Fallback to plain text extraction
            return self._extract_text_comprehensive(filepath)

    def _calculate_pdf_confidence(self, text: str, stats: Dict[str, Any]) -> float:
        """Calculate extraction confidence score."""
        if not text.strip():
            return 0.0

        base_confidence = 0.8  # Start with good base confidence

        # Adjust based on text density
        chars_per_page = len(text) / max(stats.get('pages', 1), 1)

        # Typical petrophysical report has 2000-6000 chars per page
        if 1500 <= chars_per_page <= 8000:
            base_confidence += 0.1
        elif chars_per_page < 500:
            base_confidence -= 0.3  # Might be image-based

        # Check for content quality indicators
        word_count = len(text.split())
        if word_count > 50:  # Substantial content
            base_confidence += 0.05

        return min(1.0, max(0.0, base_confidence))

    def _validate_extraction(self, text: str, metadata: Dict[str, Any]) -> List[str]:
        """Validate extraction quality and return warnings."""
        warnings = []

        if not text.strip():
            warnings.append("NO_TEXT_EXTRACTED")
            return warnings

        # Check for suspiciously low content
        if len(text) < 100 and metadata.get('file_size', 0) > 10000:
            warnings.append("LOW_TEXT_RATIO")

        # Check for excessive whitespace
        whitespace_ratio = len(re.findall(r'\s', text)) / max(len(text), 1)
        if whitespace_ratio > 0.6:
            warnings.append("HIGH_WHITESPACE_RATIO")

        # Check confidence score
        if metadata.get('confidence', 1.0) < 0.7:
            warnings.append("LOW_CONFIDENCE_SCORE")

        return warnings

    def _calculate_checksum(self, filepath: Path) -> str:
        """Calculate MD5 checksum of file."""
        hash_md5 = hashlib.md5()
        with open(str(filepath), "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def _save_results(self, results: Dict[str, ExtractionResult], output_dir: str):
        """Save extraction results to disk."""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        # Save individual text files
        text_dir = output_path / "extracted_text"
        text_dir.mkdir(exist_ok=True)

        # Save metadata
        metadata_dir = output_path / "metadata"
        metadata_dir.mkdir(exist_ok=True)

        for filename, result in results.items():
            # Save extracted text
            text_filename = text_dir / f"{Path(filename).stem}.txt"
            with open(text_filename, 'w', encoding='utf-8') as f:
                f.write(result.text)

            # Save metadata
            meta_filename = metadata_dir / f"{Path(filename).stem}.json"
            with open(meta_filename, 'w', encoding='utf-8') as f:
                json.dump({
                    'filename': result.metadata.filename,
                    'filepath': result.metadata.filepath,
                    'file_size': result.metadata.file_size,
                    'page_count': result.metadata.page_count,
                    'character_count': result.metadata.character_count,
                    'word_count': result.metadata.word_count,
                    'extraction_method': result.metadata.extraction_method,
                    'confidence_score': result.metadata.confidence_score,
                    'processing_time': result.metadata.processing_time,
                    'checksum': result.metadata.checksum,
                    'warnings': result.warnings
                }, f, indent=2)

        self.logger.info(f"Saved {len(results)} documents to {output_dir}")
